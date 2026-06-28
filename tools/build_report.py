#!/usr/bin/env python3
"""Construye el reporte técnico completo del Proyecto Final del Equipo 19."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG = DOCS / "figures"
OUTPUT = DOCS / "Proyecto_Final_Equipo19.docx"
METRICS = json.loads(
    (ROOT / "controllers/cil_autonomous_driver/training_metrics.json").read_text()
)
MANIFEST = json.loads((ROOT / "dataset/dataset_manifest.json").read_text())

BLUE = "2E74B5"
DARK_BLUE = "1F4E79"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, name: str, size: float) -> None:
    """Set a font deterministically for Word and LibreOffice renderers."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def iter_paragraphs(container):
    """Yield paragraphs in a document part, including nested table cells."""
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def enforce_requested_typography(document) -> None:
    """Arial 12 pt and black everywhere, except monospaced code snippets."""
    for style_name in (
        "Normal",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Header",
        "Footer",
        "List Bullet",
        "List Bullet 2",
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(12)
        style.font.color.rgb = BLACK
        r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{key}"), "Arial")

    containers = [document]
    for section in document.sections:
        containers.extend((section.header, section.footer))
    for container in containers:
        for paragraph in iter_paragraphs(container):
            for run in paragraph.runs:
                current = (run.font.name or "").lower()
                if current in {"menlo", "consolas", "courier new"}:
                    set_run_font(run, "Menlo", run.font.size.pt if run.font.size else 6.8)
                else:
                    set_run_font(run, "Arial", 12)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header_row(row) -> None:
    """Repeat table headers and keep them attached to the first body row."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    keep_row(row)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def add_table(doc, headers, rows, widths=None, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    repeat_header_row(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        set_cell_shading(cell, "E7E6E6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = BLACK
            run.font.size = Pt(12)
    for row_data in rows:
        cells = table.add_row().cells
        keep_row(table.rows[-1])
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if status_col is not None and i == status_col:
                text = str(value).lower()
                fill = GREEN if any(k in text for k in ("cumple", "verificado", "aprobado")) else AMBER
                set_cell_shading(cells[i], fill)
            for para in cells[i].paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(12)
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, "Arial", 12)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = BLACK


def add_figure(doc, path: Path, caption: str, width=6.4) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_bullet(doc, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)


def add_note(doc, title: str, body: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + ". ")
    r.bold = True
    r.font.color.rgb = BLACK
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, title: str, code: str, font_size=6.8, chunk_size=32) -> None:
    doc.add_heading(title, level=2)
    lines = code.splitlines()
    for start in range(0, len(lines), chunk_size):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_together = False
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F7F7F7")
        p_pr.append(shd)
        r = p.add_run("\n".join(lines[start : start + chunk_size]))
        set_run_font(r, "Menlo", font_size)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)
for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    style = styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    style.font.color.rgb = BLACK
    style.font.bold = True
    style.paragraph_format.space_before = Pt(14 if name != "Title" else 0)
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.keep_with_next = True

doc.core_properties.title = "Proyecto Final de Navegación Autónoma — Equipo 19"
doc.core_properties.author = "Equipo 19"
doc.core_properties.subject = "Conditional Imitation Learning y seguridad en Webots R2025a"

# Portada basada en Plantilla-Tareas-TEC.docx
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
p.add_run().add_picture(str(FIG / "tec_logo.png"), width=Inches(4.7))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Escuela de Ingeniería y Ciencias")
r.font.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Maestría en Inteligencia Artificial Aplicada")
r.font.size = Pt(16)
r.font.bold = True
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("PRESENTA:").bold = True
for name, student_id in (
    ("José Antonio Fernández Pérez", "A01795991"),
    ("Deménard Gardy Armand", "A01797139"),
    ("Luis Daniel Castillo Alegría", "A01224696"),
    ("Raúl Adrián Delgado Rodríguez", "A01246414"),
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{name} · {student_id}")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("ACTIVIDAD:\n").bold = True
p.add_run("Proyecto Final: Navegación Autónoma con Conditional Imitation Learning\n").bold = True
p.add_run("Equipo 19\n\n").bold = True
p.add_run("Curso: Navegación autónoma\n")
p.add_run("Profesor: Dr. David Antonio Torres\n\n")
p.add_run("30 de junio de 2026").bold = True
doc.add_page_break()

doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph(
    "Se implementó una solución de Conditional Imitation Learning (CIL) con tres comandos de alto nivel: "
    "STRAIGHT=0, LEFT=1 y RIGHT=2. El flujo corrige el conjunto original sin encabezado, excluye 321 muestras "
    "de recuperación, impide fuga entre entrenamiento y validación mediante una división por fuente y aumenta "
    "el conjunto hasta 12,956 muestras. La CNN condicionada alcanzó un MAE de validación global de 0.0220 rad, "
    "inferior al objetivo de 0.05 rad."
)
doc.add_paragraph(
    "En Webots R2025a se integraron Camera Recognition, LiDAR, radar, giroscopio y tres sensores laterales. "
    "El árbitro de seguridad prioriza el frenado ante peatones, la evasión del autobús estacionado, el control "
    "de distancia y, finalmente, la dirección predicha por CIL. Se verificaron dos recorridos limpios por ruta, "
    "incluyendo evasión, frenado ante peatón y seguimiento de tráfico."
)
add_note(
    doc,
    "Estado de cierre",
    "El reporte técnico, el código, el modelo, el notebook y la validación completa de las tres rutas están "
    "cerrados. Sólo faltan la edición con participación de los cuatro integrantes y la publicación del video; "
    "por integridad académica no se inventa una URL.",
    fill=AMBER,
)

doc.add_heading("1. Planteamiento y objetivos", level=1)
doc.add_paragraph(
    "El objetivo es conducir el vehículo autónomo por tres rutas urbanas mediante imitación condicional, "
    "manteniendo velocidad de evaluación de 22 km/h y un límite absoluto de 30 km/h. La solución debe reaccionar "
    "ante vehículos próximos, peatones y un autobús estacionado, además de mantener una ejecución reproducible."
)
doc.add_paragraph(
    "Codevilla et al. propusieron condicionar la política de conducción con un comando de navegación de alto nivel. "
    "Su planteamiento original contempla Follow lane, Left, Right y Straight. Para la evaluación de este proyecto "
    "se adapta explícitamente el espacio de comandos a las tres rutas solicitadas y se elimina la etiqueta recovery; "
    "la recuperación ante obstáculos es una máquina de estados de seguridad, no una clase aprendida."
)
add_bullet(doc, "Ruta recta: torres residenciales del noreste → silos; tráfico y evasión del autobús.")
add_bullet(doc, "Ruta derecha: Subway norte → iglesia; detección y frenado ante peatón.")
add_bullet(doc, "Ruta izquierda: parque infantil → Subway norte.")

doc.add_heading("2. Preparación e integridad del conjunto de datos", level=1)
doc.add_paragraph(
    "El CSV original carecía de encabezado y contenía cuatro comandos. El proceso de saneamiento conserva sólo "
    "0, 1 y 2, mueve las 321 filas recovery a un archivo histórico no publicado y comprueba tanto referencias "
    "faltantes como imágenes huérfanas. El conjunto canónico contiene 6,129 imágenes PNG."
)
add_table(
    doc,
    ["Comando", "Etiqueta", "Muestras originales", "Muestras de entrenamiento aumentadas", "Validación"],
    [
        (0, "Straight", 1202, 3848, 240),
        (1, "Left", 1895, 3941, 379),
        (2, "Right", 3032, 3941, 607),
        ("Total", "", 6129, 11730, 1226),
    ],
    widths=[0.65, 1.0, 1.35, 1.75, 1.0],
)
add_figure(doc, FIG / "command_distribution_original.png", "Figura 1. Distribución real de comandos antes y después del aumento.")
doc.add_heading("2.1 División sin fuga y aumento", level=2)
doc.add_paragraph(
    "La división estratificada 80/20 se realiza sobre `source_id` antes de cualquier transformación. Sólo después "
    "se generan reflejos horizontales —intercambiando LEFT y RIGHT— y variaciones de brillo para reforzar STRAIGHT. "
    "La intersección de fuentes entre entrenamiento y validación es cero."
)
add_table(
    doc,
    ["Comprobación", "Resultado", "Estado"],
    [
        ("Imágenes originales faltantes", 0, "Cumple"),
        ("Imágenes originales no referenciadas", 0, "Cumple"),
        ("Comandos válidos", "0, 1, 2", "Cumple"),
        ("Muestras después de aumento", 12956, "Cumple > 10,000"),
        ("Fuentes compartidas train/val", 0, "Cumple"),
    ],
    widths=[3.2, 1.4, 1.4],
    status_col=2,
)

doc.add_heading("3. Arquitectura y entrenamiento", level=1)
doc.add_paragraph(
    "La red recibe una imagen RGB preprocesada de 80×160×3 y un vector one-hot de tres posiciones. Cinco capas "
    "convolucionales extraen características; la rama visual se aplana y concatena con el comando. Capas densas "
    "producen un único ángulo de dirección en radianes. Se emplearon semilla reproducible, Adam, pérdida MSE, MAE, "
    "EarlyStopping, ReduceLROnPlateau y checkpoint del mejor HDF5."
)
add_figure(doc, FIG / "model_architecture.png", "Figura 2. Arquitectura obtenida directamente del modelo HDF5 entrenado.", width=5.0)
add_table(
    doc,
    ["Parámetro", "Valor"],
    [
        ("Tamaño de entrada", "80 × 160 × 3"),
        ("Comando", "One-hot de 3 clases"),
        ("Optimizador", "Adam"),
        ("Pérdida", "MSE"),
        ("Callbacks", "EarlyStopping, ReduceLROnPlateau, ModelCheckpoint"),
        ("Formato de salida", "HDF5 + model_metadata.json"),
    ],
    widths=[2.3, 4.0],
)
add_figure(doc, FIG / "training_curves.png", "Figura 3. Curvas reales de pérdida y MAE durante el entrenamiento.")
doc.add_heading("3.1 Resultados cuantitativos", level=2)
add_table(
    doc,
    ["Métrica", "Resultado (rad)", "Criterio"],
    [
        ("MAE global de validación", f"{METRICS['validation_mae_rad']:.4f}", "≤ 0.05 — cumple"),
        ("MAE Straight", f"{METRICS['per_command'][0]['mae_rad']:.4f}", "Reportado"),
        ("MAE Left", f"{METRICS['per_command'][1]['mae_rad']:.4f}", "Reportado"),
        ("MAE Right", f"{METRICS['per_command'][2]['mae_rad']:.4f}", "Reportado"),
        ("MSE global", f"{METRICS['validation_mse']:.4f}", "Reportado"),
    ],
    widths=[2.8, 1.5, 2.0],
)

doc.add_heading("4. Controladores y portabilidad", level=1)
doc.add_heading("4.1 Captura en World 1", level=2)
doc.add_paragraph(
    "El controlador de captura usa rutas relativas al proyecto, crea el encabezado `image,steering,command`, "
    "limita la velocidad a 30 km/h y registra cada 100 ms. Las teclas W, A y D asignan los tres comandos válidos."
)
doc.add_heading("4.2 Control autónomo en World 2", level=2)
doc.add_paragraph(
    "El controlador carga `cil_model.h5` y `model_metadata.json` desde su propio directorio, replica el "
    "preprocesamiento del notebook, limita ángulo y tasa de cambio, y ejecuta inferencia cada cuatro pasos. "
    "El lanzador regenera presets relativos, crea el entorno Python, localiza SUMO y conecta el controlador externo."
)
add_table(
    doc,
    ["Prioridad", "Condición", "Acción"],
    [
        (1, "Peatón ≤ 15 m por Recognition y LiDAR", "Freno total"),
        (2, "Autobús estacionado ≤ 18 m", "Separación, pared derecha, orientación y reincorporación"),
        (3, "Vehículo entre 25 y 12 m", "Regulación progresiva; alto ≤12 m; reanuda ≥15 m"),
        (4, "Sin riesgo superior", "Dirección CIL; 22 km/h"),
    ],
    widths=[0.7, 2.6, 3.0],
)
doc.add_heading("4.3 Máquina de estados de evasión", level=2)
add_table(
    doc,
    ["Estado", "Propósito", "Criterio de transición"],
    [
        ("CIL", "Conducción nominal", "Autobús reconocido a ≤18 m"),
        ("EVASION_SEPARACION", "Desplazamiento inicial a la izquierda", "Pared lateral o 3.5 s"),
        ("SEGUIMIENTO_PARED_DERECHA", "Mantener separación del autobús", "Pared ausente durante 12 pasos"),
        ("RECUPERA_ORIENTACION", "Compensar giro con giroscopio", "Error <0.08 rad o 6 s"),
        ("REINCORPORACION", "Volver suavemente a CIL", "2 s"),
    ],
    widths=[1.7, 2.2, 2.4],
)

doc.add_heading("5. Mundo, sensores y tráfico", level=1)
doc.add_paragraph(
    "World 2 incorpora Camera Recognition, Sick LMS 291, radar frontal, giroscopio y tres sensores de distancia "
    "en el costado derecho. `RecognizablePedestrian.proto` es local, tiene modelo `pedestrian`, colores de "
    "reconocimiento y objetos de colisión. `SumoInterface` limita explícitamente la simulación a 30 vehículos."
)
add_table(
    doc,
    ["Dispositivo", "Uso"],
    [
        ("Camera + Recognition", "Clasificar peatón/autobús y estimar distancia relativa"),
        ("LiDAR", "Confirmar obstáculo frontal en el campo de la cámara"),
        ("Radar", "Distancia y velocidad radial del vehículo precedente"),
        ("Giroscopio", "Recuperar orientación tras evasión"),
        ("3 sensores laterales", "Seguimiento de pared derecha y detección de fin de obstáculo"),
    ],
    widths=[2.1, 4.2],
)

doc.add_heading("6. Presets de ruta y evidencia de Webots", level=1)
doc.add_paragraph(
    "Los presets se regeneran mediante `tools/create_route_worlds.py` y comparten la misma red vial. Cada uno "
    "fija pose inicial y comando, conservando rutas relativas. Las figuras provienen directamente de la cámara "
    "del vehículo durante Webots R2025a."
)
add_figure(doc, FIG / "webots_straight.png", "Figura 4. Ruta recta: evidencia final de llegada al corredor de los silos.")
add_figure(doc, FIG / "webots_right.png", "Figura 5. Ruta derecha: evidencia final tras freno ante peatón y giro.")
add_figure(doc, FIG / "webots_left.png", "Figura 6. Ruta izquierda: evidencia final en el corredor de Subway norte.")
add_table(
    doc,
    ["Ruta", "Comando", "Resultado técnico observado"],
    [
        ("Recto", 0, "Evasión completa; llegadas GPS=(-188.08,235.10) y (-188.03,235.40)"),
        ("Derecha", 2, "Freno a peatón; llegadas GPS=(23.96,-89.60) y (28.97,-65.41)"),
        ("Izquierda", 1, "Llegadas GPS=(35.04,25.75) y (35.04,28.37)"),
    ],
    widths=[1.2, 1.0, 4.1],
)
add_note(
    doc,
    "Validación de rutas cerrada",
    "Se ejecutaron dos recorridos limpios por ruta, por carril derecho, sin colisión ni U-turn. Los logs registran "
    "los cambios de estado y la llegada GPS. La validación también corrigió la falsa detección de 'bus stop' "
    "como autobús mediante filtrado de modelo y prueba de regresión.",
    fill=GREEN,
)

doc.add_heading("7. Verificación automatizada", level=1)
doc.add_paragraph(
    "Doce pruebas unitarias y de integridad verifican esquema y referencias del dataset, notebook reproducible, "
    "presets, dispositivos del mundo, preprocesamiento, one-hot, límites de dirección, regulación longitudinal, "
    "histéresis de 12/15 m y acotación del seguimiento de pared."
)
add_table(
    doc,
    ["Área", "Evidencia", "Estado"],
    [
        ("Dataset", "6,129 imágenes; 0 faltantes; 3 comandos", "Verificado"),
        ("Aumento", "12,956 muestras; 0 source_id compartidos", "Verificado"),
        ("Notebook", "19 celdas ejecutadas desde clon; HDF5 recargable", "Verificado"),
        ("Modelo", "MAE 0.0220 rad y métricas por comando", "Verificado"),
        ("Pruebas", "12/12 exitosas", "Verificado"),
        ("Webots", "Tres comandos y árbitro de seguridad", "Verificado"),
        ("Recorridos completos", "Dos tomas limpias por ruta", "Verificado"),
        ("YouTube", "Video <6 min y URL pública", "Pendiente de publicación"),
    ],
    widths=[1.5, 3.3, 1.5],
    status_col=2,
)

doc.add_heading("8. Reproducción", level=1)
doc.add_paragraph("Entrenamiento: abrir el notebook en Colab y ejecutar todas las celdas. La primera celda clona el repositorio público.")
add_code(doc, "Comandos de aceptación local", """git clone https://github.com/xak47d/proyecto-final-cil-equipo19.git
cd proyecto-final-cil-equipo19
./run_final_project.sh --setup
./run_final_project.sh --route straight --record
./run_final_project.sh --route right --record
./run_final_project.sh --route left --record
.venv-final/bin/python -m unittest discover -s tests -v""")

doc.add_heading("9. Uso de inteligencia artificial", level=1)
doc.add_paragraph(
    "Se utilizó OpenAI Codex basado en GPT-5 como asistente para generación de código, depuración, optimización, "
    "estructuración del notebook y del reporte, y preparación de pruebas. El equipo conserva responsabilidad sobre "
    "la revisión técnica, la ejecución en Webots, la interpretación de resultados y la entrega académica."
)

doc.add_heading("10. Conclusiones", level=1)
doc.add_paragraph(
    "La solución elimina correctamente la clase recovery y separa el aprendizaje de dirección de la recuperación "
    "operativa. El modelo supera la meta cuantitativa, el dataset queda trazable y libre de fuga, y el controlador "
    "integra sensores heterogéneos con prioridades verificables y completó dos recorridos limpios por ruta. Los "
    "únicos pasos externos pendientes son la edición con los cuatro integrantes y la publicación del video."
)

doc.add_heading("11. Matriz de cumplimiento de la rúbrica", level=1)
matrix_rows = [
    ("Dataset y comandos", "CSV con encabezado; 321 recovery excluidas; 0/1/2", "Cumple"),
    ("Aumento y división", ">10,000; split por fuente; reflejo y brillo", "Cumple"),
    ("Notebook", "Clon, semillas, distribución, arquitectura, curvas, MAE y exportación", "Cumple"),
    ("Modelo CIL", "CNN condicionada + callbacks", "Cumple"),
    ("Portabilidad", "Rutas relativas y lanzador reproducible", "Cumple"),
    ("Sensores", "Recognition, LiDAR, radar, gyro, 3 laterales", "Cumple"),
    ("Peatón", "≤15 m, cámara + LiDAR, freno total", "Cumple"),
    ("Control de distancia", "25–12 m; alto ≤12; reanuda ≥15", "Cumple"),
    ("Autobús", "≤18 m + FSM de evasión", "Cumple"),
    ("SUMO", "maxVehicles 30", "Cumple"),
    ("Tres rutas", "Presets, logs y clips completos", "Cumple"),
    ("Recorridos completos", "Dos recorridos limpios por ruta", "Cumple"),
    ("Reporte", "Español, anexos completos, figuras reales", "Cumple"),
    ("Video", "Guion 5:35, cuatro integrantes", "Guion cumple; grabación pendiente"),
    ("YouTube", "URL pública real", "Pendiente de publicación"),
]
add_table(doc, ["Criterio", "Evidencia", "Estado"], matrix_rows, widths=[1.7, 3.2, 1.4], status_col=2)

doc.add_heading("Referencias", level=1)
refs = [
    "Codevilla, F., Müller, M., López, A., Koltun, V. y Dosovitskiy, A. (2018). End-to-End Driving Via Conditional Imitation Learning. https://arxiv.org/abs/1710.02410",
    "CARLA Simulator. Imitation Learning — repositorio oficial. https://github.com/carla-simulator/imitation-learning",
    "Bojarski, M. et al. (2016). End to End Learning for Self-Driving Cars. arXiv:1604.07316.",
    "Cyberbotics. (2025). Webots R2025a Reference Manual. https://cyberbotics.com/doc/reference/index",
    "Eclipse SUMO. (2026). Simulation of Urban MObility documentation. https://sumo.dlr.de/docs/",
    "OpenAI. (2026). Codex (basado en GPT-5) [asistente de programación].",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)

# Anexos en orientación horizontal para preservar código íntegro.
land = doc.add_section(WD_SECTION.NEW_PAGE)
land.orientation = WD_ORIENT.LANDSCAPE
land.page_width = Inches(11)
land.page_height = Inches(8.5)
land.top_margin = Inches(0.55)
land.bottom_margin = Inches(0.55)
land.left_margin = Inches(0.55)
land.right_margin = Inches(0.55)

doc.add_heading("Anexo A. Controlador de captura completo", level=1)
doc.add_paragraph("Código íntegro y comentado de `controllers/cil_dataset_recorder/cil_dataset_recorder.py`.")
add_code(doc, "A.1 Fuente", (ROOT / "controllers/cil_dataset_recorder/cil_dataset_recorder.py").read_text())
doc.add_page_break()
doc.add_heading("Anexo B. Controlador autónomo completo", level=1)
doc.add_paragraph("Código íntegro y comentado de `controllers/cil_autonomous_driver/cil_autonomous_driver.py`.")
add_code(doc, "B.1 Fuente", (ROOT / "controllers/cil_autonomous_driver/cil_autonomous_driver.py").read_text())
doc.add_heading("Anexo C. Celdas de código del notebook", level=1)
nb = json.loads((ROOT / "Proyecto_Final_Equipo19.ipynb").read_text())
code_cells = [cell for cell in nb["cells"] if cell.get("cell_type") == "code"]
for index, cell in enumerate(code_cells, 1):
    source = "".join(cell.get("source", []))
    add_code(doc, f"C.{index} Celda {index}", source)
doc.add_heading("Anexo D. Guion del video", level=1)
doc.add_paragraph(
    "El guion dura aproximadamente 5:35, distribuye la participación entre los cuatro integrantes y evita "
    "superar los seis minutos. La URL de YouTube se incorporará sólo después de su publicación."
)
add_code(
    doc,
    "D.1 Guion completo",
    (DOCS / "Guion_Video_Equipo19.md").read_text(),
    font_size=5.8,
    chunk_size=100,
)

# La entrega no lleva encabezados, pies de página ni numeración.
for section in doc.sections:
    for part in (
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
        section.even_page_header,
        section.even_page_footer,
    ):
        part.is_linked_to_previous = False
        for paragraph in part.paragraphs:
            for child in list(paragraph._p):
                paragraph._p.remove(child)

enforce_requested_typography(doc)
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
